import streamlit as st
import pandas as pd
import json, os, base64, requests, re
from io import BytesIO
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.drawing.image import Image as XLImage
import matplotlib.pyplot as plt
import hashlib

# AI
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.svm import LinearSVC
from sklearn.calibration import CalibratedClassifierCV
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION_START


# =============== CONFIG =================
st.set_page_config(page_title="Studio ISA e Registro IVA", layout="wide")

GITHUB_FILE_A = "dizionario_drveto_stage.json"
GITHUB_FILE_B = "dizionario_vetsgo_stage.json"
CONFIG_FILE = "config_clinica_stage.json"
USERS_FILE = "users.json"
GITHUB_REPO = os.getenv("GITHUB_REPO")
GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")

# AI MODELS (global placeholders)
model = None
vectorizer = None
model_B = None
vectorizer_B = None

# =============== UTILS =================
def norm(s):
    return re.sub(r"\s+", " ", str(s).strip().lower())


def any_kw_in(t, keys):
    return any(k in t for k in keys)


def coerce_numeric(s):
    if pd.api.types.is_numeric_dtype(s):
        return pd.to_numeric(s, errors="coerce").fillna(0)
    s = (
        s.astype(str)
        .str.replace(r"\s", "", regex=True)
        .str.replace("€", "", regex=False)
        .str.replace(".", "", regex=False)
        .str.replace(",", ".", regex=False)
    )
    return pd.to_numeric(s, errors="coerce").fillna(0)


def round_pct(values):
    values = pd.to_numeric(values, errors="coerce").fillna(0)
    total = Decimal(str(values.sum()))
    if total == 0:
        return values * 0
    raw = [Decimal(str(v)) * Decimal("100") / total for v in values]
    rounded = [r.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP) for r in raw]
    diff = Decimal("100.00") - sum(rounded)
    rounded[-1] = (rounded[-1] + diff).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return pd.Series([float(x) for x in rounded], index=values.index)


@st.cache_data(ttl=600, show_spinner=False)
def load_excel(file):
    return pd.read_excel(file)


# =============== GITHUB =================
def load_users():
    data = github_load_json(USERS_FILE)
    for k, v in data.items():
        if "permissions" not in v:
            v["permissions"] = {
                "manage_ai": False,
                "use_registro_iva": True,
                "manage_clinics": False,
                "manage_users": False,
            }
        if "clinics" not in v:
            v["clinics"] = []
    return data

def save_users(users: dict):
    github_save_json(USERS_FILE, users)

def reset_admin_password(new_password):
    users = load_users()
    if "admin" not in users:
        users["admin"] = {"password": hash_pwd(new_password), "role": "admin", "permissions": {}}
    else:
        users["admin"]["password"] = hash_pwd(new_password)
    save_users(users)

def hash_pwd(pwd: str) -> str:
    return hashlib.sha256(pwd.encode()).hexdigest()

def github_load_json(file_name):
    try:
        url = f"https://api.github.com/repos/{GITHUB_REPO}/contents/{file_name}"
        headers = {"Authorization": f"token {GITHUB_TOKEN}"} if GITHUB_TOKEN else {}
        r = requests.get(url, headers=headers, timeout=12)
        if r.status_code == 200:
            raw = json.loads(base64.b64decode(r.json()["content"]).decode("utf-8"))
            return raw
    except:
        pass
    return {}


def github_save_json(file_name, data):
    url = f"https://api.github.com/repos/{GITHUB_REPO}/contents/{file_name}"
    headers = {"Authorization": f"token {GITHUB_TOKEN}"}
    
    get = requests.get(url, headers=headers)
    sha = get.json().get("sha") if get.status_code == 200 else None
    
    encoded = base64.b64encode(json.dumps(data, ensure_ascii=False, indent=2).encode()).decode()
    payload = {"message": "Update ISA dictionary", "content": encoded, "branch": "main"}
    if sha:
        payload["sha"] = sha
    
    r = requests.put(url, headers=headers, data=json.dumps(payload))
    if r.status_code not in (200, 201):
        st.error(f"❌ Errore salvataggio GitHub: {r.status_code} → {r.text}")

# =============== LOGIN SYSTEM =================
def login():
    users = load_users()

    if "logged_user" in st.session_state:
        return st.session_state.logged_user  # già loggato

    st.title("🔐 Accesso al portale")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Accedi"):
        if username in users and users[username]["password"] == hash_pwd(password):
            st.session_state.logged_user = username

            #aggiorna stato ultimo
            users[username]["last_login"] = datetime.now().isoformat()
            save_users(users)
            
            st.success(f"Benvenuto {username} 👋")
            st.rerun()
        else:
            st.error("❌ Credenziali errate")
    
    st.stop()  # blocca l’app finché non fa login

logged_user = login()

def load_clinic_config():
    try:
        url = f"https://api.github.com/repos/{GITHUB_REPO}/contents/{CONFIG_FILE}"
        headers = {"Authorization": f"token {GITHUB_TOKEN}"} if GITHUB_TOKEN else {}
        r = requests.get(url, headers=headers, timeout=12)
        if r.status_code == 200:
            raw = base64.b64decode(r.json()["content"]).decode("utf-8")
            return json.loads(raw)
    except:
        pass
    return {}

def save_clinic_config(data: dict):
    try:
        url = f"https://api.github.com/repos/{GITHUB_REPO}/contents/{CONFIG_FILE}"
        headers = {"Authorization": f"token {GITHUB_TOKEN}"}
        r = requests.get(url, headers=headers, timeout=12)
        sha = r.json().get("sha") if r.status_code == 200 else None

        encoded = base64.b64encode(json.dumps(data, ensure_ascii=False, indent=2).encode()).decode()
        payload = {"message": "Update clinic config", "content": encoded, "branch": "main"}
        if sha:
            payload["sha"] = sha

        requests.put(url, headers=headers, data=json.dumps(payload), timeout=20)
    except:
        pass

# =============== CATEGORY RULES =================
RULES_A = {
    "Visite domiciliari o presso allevamenti": ["domicilio","allevamenti"],
    "Visite ambulatoriali": ["visita","controllo","consulto","terapia","trattam","manual","microchip","vacc","medicazione"],
    "Esami diagnostici per immagine": ["rx","eco","ecogra","tac","raggi","radi"],
    "Altri esami diagnostici": ["analisi","emocromo","prelievo","laboratorio"],
    "Interventi chirurgici": ["chirurg","castraz","ovariect","detartrasi","estraz","anest","endo"],
    "Altre attività": ["acconto"]
}

RULES_B = {
    "Visite domiciliari o presso allevamenti": ["domicilio","allevamenti"],
    "Visite ambulatoriali": ["visita","controllo","consulto","terapia","trattam","manual","microchip","vacc","medicazione"],
    "Esami diagnostici per immagine": ["rx","eco","ecogra","tac","raggi","radi"],
    "Altri esami diagnostici": ["analisi","emocromo","prelievo","laboratorio"],
    "Interventi chirurgici": ["chirurg","castraz","ovariect","detartrasi","estraz","anest","endo"],
    "Altre attività": ["acconto"]
}

ORDER_B = list(RULES_B.keys()) + ["Totale"]


# =============== AI TRAIN =================
def train_ai_model(dictionary):
    dictionary = {k: v for k, v in dictionary.items() if k.strip() != ""}
    if not dictionary:
        return None, None

    texts = list(dictionary.keys())
    labels = list(dictionary.values())

    vec = TfidfVectorizer(lowercase=True, ngram_range=(1,2), min_df=1)
    try:
        X = vec.fit_transform(texts)
    except ValueError:
        return None, None

    if len(set(labels)) <= 1:
        return None, None

    # SVM molto più precisa
    base_model = LinearSVC()
    model = CalibratedClassifierCV(base_model)  # aggiunge predict_proba() per auto-apprendimento
    model.fit(X, labels)

    return vec, model

# =============== CLASSIFICATION (helpers) =================
def classify_A(desc, fam, mem):
    """Rule-based + memory + (optional) AI suggestion (used only in auto-pass)."""
    global model, vectorizer
    d = norm(desc)

    fam_s = norm(fam)
    if fam_s and fam_s not in {"privato","professionista","nan","none",""}:
        return fam_s.upper()
    d_tokens = d.split()
    for k, v in mem.items():
        k_norm = norm(k)
        if k_norm == d or k_norm in d_tokens:
            return v

    # Pure rule fallback
    for cat, keys in RULES_A.items():
        if any_kw_in(d, keys):
            return cat

    return "ALTRE PRESTAZIONI"


def classify_B(prest, mem):
    """Rule-based + memory + (optional) AI suggestion (used only in auto-pass)."""
    d = norm(prest)

    d_tokens = d.split()
    for k, v in mem.items():
        k_norm = norm(k)
        if k_norm == d or k_norm in d_tokens:
            return v

    for cat, keys in RULES_B.items():
        if any_kw_in(d, keys):
            return cat

    return "Altre attività"

def render_user_management():
    st.title("👤 Gestione Utenti")

    users = load_users()

    # --- Ricerca utenti ---
    search = st.text_input("🔍 Cerca utente", placeholder="Inserisci nome utente...").strip().lower()
    filtered = {u: d for u, d in users.items() if search in u.lower()}

    # --- Tabella utenti ---
    st.subheader(f"Utenti trovati: {len(filtered)}")

    if not filtered:
        st.info("Nessun utente trovato.")
        return

    df_users = pd.DataFrame([
        {
            "Username": u,
            "Ruolo": d.get("role", "user"),
            "Gestione AI": "✅" if d.get("permissions", {}).get("manage_ai", False) else "❌",
            "Gestione Cliniche (aggiungi/rimuovi)": "✅" if d.get("permissions", {}).get("manage_clinics", False) else "❌",
            "Può usare Registro IVA": "✅" if d.get("permissions", {}).get("use_registro_iva", False) else "❌",
            "Gestione Utenti": "✅" if d.get("permissions", {}).get("manage_users", False) else "❌",
        }
        for u, d in filtered.items()
    ])

    st.dataframe(df_users, use_container_width=True)

    st.markdown("---")
    st.subheader("✏️ Modifica Utente")

    selected = st.selectbox("Seleziona utente", list(filtered.keys()))
    u = users[selected]

    col1, col2 = st.columns(2)

    with col1:
        role_new = st.selectbox(
            "Ruolo",
            ["user", "admin"],
            index=["user", "admin"].index(u.get("role", "user")),
            key=f"role_{selected}"
        )

        # --- PERMESSI DI BASE ---
        perms = u.get("permissions", {})

        p1 = st.checkbox("Può modificare sensibilità AI", value=perms.get("manage_ai", False), key=f"p_ai_{selected}")
        p2 = st.checkbox("Può usare Registro IVA", value=perms.get("use_registro_iva", True), key=f"p_registro_{selected}")
        p3 = st.checkbox("Può gestire Cliniche (aggiungi/modifica)", value=perms.get("manage_clinics", False), key=f"p_clinic_{selected}")
        p4 = st.checkbox("Può gestire utenti", value=perms.get("manage_users", False), key=f"p_users_{selected}")
        
        st.markdown("### 🏥 Cliniche assegnate a questo utente:")

        config_all = load_clinic_config()
        cliniche_disponibili = list(config_all.keys())
        cliniche_utente = set(u.get("clinics", []))
        cliniche_nuove = set()

        for c in cliniche_disponibili:
            checked = st.checkbox(c, value=(c in cliniche_utente), key=f"clinic_{selected}_{c}")
            if checked:
                cliniche_nuove.add(c)

        if st.button("💾 Salva modifiche", key=f"save_all_{selected}"):
            users[selected]["role"] = role_new
            users[selected]["permissions"] = {
                "manage_ai": p1,
                "use_registro_iva": p2,
                "manage_clinics": p3,
                "manage_users": p4,
            }
            users[selected]["clinics"] = sorted(list(cliniche_nuove))
            save_users(users)
            st.success("✅ Permessi e cliniche aggiornate.")
            st.rerun()

    with col2:
        new_pwd = st.text_input("Nuova password", type="password", placeholder="Lascia vuoto per non cambiare")
        if st.button("🔑 Reset Password"):
            if new_pwd.strip():
                users[selected]["password"] = hash_pwd(new_pwd)
                save_users(users)
                st.success("✅ Password aggiornata.")
                st.rerun()
            else:
                st.warning("⚠️ Inserisci una password valida.")

        if selected not in ("admin", logged_user):
            if st.button("🗑️ Elimina Utente"):
                users.pop(selected)
                save_users(users)
                st.success("✅ Utente eliminato.")
                st.rerun()

    st.markdown("---")
    st.subheader("➕ Crea Nuovo Utente")
    new_user = st.text_input("Username nuovo utente")
    new_pwd2 = st.text_input("Password", type="password")
    role_new2 = st.selectbox("Ruolo nuovo", ["user", "admin"])

    if st.button("✅ Crea Utente"):
        if new_user.strip() and new_pwd2.strip():
            users[new_user] = {
                "password": hash_pwd(new_pwd2),
                "role": role_new2,
                "permissions": {"manage_ai": False, "manage_clinics": False, "manage_users": False}
            }
            save_users(users)
            st.success(f"✅ Utente {new_user} creato.")
            st.rerun()
        else:
            st.warning("⚠️ Compila tutti i campi.")

# =============== SIDEBAR =================
pages = ["📊 Studio ISA", "📄 Registro IVA"]
user_data = load_users().get(logged_user,{})
permissions = user_data.get("permissions", {})

if user_data.get("role") == "admin" or permissions.get("manage_users", False):
    pages.append("👤 Gestione Utenti")
page = st.sidebar.radio("📌 Navigazione", pages)

# --- Permessi utente ---
users_all = load_users()
user_data = users_all.get(logged_user, {})
permissions = user_data.get("permissions", {})

can_manage_ai = permissions.get("manage_ai", False)

auto_thresh = st.sidebar.slider(
    "Soglia auto-apprendimento (AI)",
    0.50, 0.99, 0.85, 0.01,
    disabled=not can_manage_ai
)

if not can_manage_ai:
    st.sidebar.caption("🔒 Non hai il permesso di modificare la sensibilità AI")
else:
    st.sidebar.caption("✅ Puoi modificare la sensibilità AI")
st.sidebar.caption("Se la confidenza del modello ≥ soglia, il termine viene appreso in automatico.")
st.sidebar.caption("Alcyon Italia SpA - 2025")
st.sidebar.caption("v.1.1)

# =============== MAIN =================
def main():
    if page == "👤 Gestione Utenti":
        render_user_management()
        st.stop()
    
    if st.sidebar.button("🔓 Logout"):
        st.session_state.pop("logged_user", None)
        st.rerun()
    global model, vectorizer, model_B, vectorizer_B

    st.title("📊 Studio ISA – DrVeto e VetsGo")

    file = st.file_uploader("Seleziona Excel", type=["xlsx","xls"])
    if not file:
        st.stop()

    # Load file only once
    if "df" not in st.session_state:
        df = load_excel(file)
        st.session_state.df = df
        mode = "B" if any("prestazioneprodotto" in c.replace(" ","").lower() for c in df.columns) else "A"
        st.session_state.mode = mode
        st.session_state.mem = github_load_json(GITHUB_FILE_A if mode=="A" else GITHUB_FILE_B)
        st.session_state.new = {}
        st.session_state.idx = 0
        st.session_state.auto_added = []  # [(term, cat, conf)]

    df = st.session_state.df.copy()
    mem = st.session_state.mem
    new = st.session_state.new
    mode = st.session_state.mode

    # Train AI
    if mode == "A":
        vectorizer, model = train_ai_model(mem | new)
    else:
        vectorizer_B, model_B = train_ai_model(mem | new)


    # --- RILEVA COLONNE UNA VOLTA SOLA ---
    if mode == "A":
        desc = next(c for c in df.columns if "descrizione" in c.lower())
        fam  = next((c for c in df.columns if "famiglia" in c.lower()), None)
        qta  = next(c for c in df.columns if "quant" in c.lower() or c.strip() == "%")
        netto = next(c for c in df.columns if "netto" in c.lower() and "dopo" in c.lower())

    else:
        prest = next(c for c in df.columns if "prestazioneprodotto" in c.replace(" ", "").lower())
        imp   = next(c for c in df.columns if "totaleimpon" in c.lower())
        iva_col = next((c for c in df.columns if "totaleconiva" in c.replace(" ", "").lower()), None)
        tot   = next(c for c in df.columns if c.lower().strip()=="totale" or "totale" in c.lower())
    
    # ===== PROCESS A =====
    if mode == "A":
        desc = next(c for c in df.columns if "descrizione" in c.lower())
        fam = next((c for c in df.columns if "famiglia" in c.lower()), None)
        qta = next(c for c in df.columns if "quant" in c.lower() or c.strip()=="%")
        netto = next(c for c in df.columns if "netto" in c.lower() and "dopo" in c.lower())

        df[qta] = coerce_numeric(df[qta])
        df[netto] = coerce_numeric(df[netto])

        base = desc

        # ========== AUTO APPRENDIMENTO PASS ==========
        learned = {norm(k) for k in (mem | new).keys()}
        df["_clean"] = df[base].astype(str).map(norm)
        candidates = sorted([t for t in df["_clean"].unique() if t not in learned])

        auto_added_now = []
        if model and vectorizer and candidates:
            X = vectorizer.transform(candidates)
            probs = model.predict_proba(X)
            preds = model.classes_[probs.argmax(axis=1)]
            confs = probs.max(axis=1)
            for t, p, c in zip(candidates, preds, confs):
                SAFE_CONFIDENCE = max(auto_thresh, 0.95)
                if float(c) >= SAFE_CONFIDENCE:
                    auto_added_now.append((t, p, float(c)))

        if auto_added_now:
            st.session_state.new = new
            st.session_state.auto_added.extend(auto_added_now)
            mem.update(new)
            github_save_json(GITHUB_FILE_A if mode=="A" else GITHUB_FILE_B, mem)
            st.session_state.mem = mem
            st.info(f"✅ Salvato sul cloud ({len(auto_added_now)} nuovi termini appresi automaticamente)")

        # Classify rows (using classify_A that includes memory & rules)
        df["CategoriaFinale"] = df.apply(lambda r: classify_A(r[desc], r[fam] if fam else None, mem | new), axis=1)

    # ===== PROCESS B =====
    else:
        prest = next(c for c in df.columns if "prestazioneprodotto" in c.replace(" ","").lower())
        imp = next(c for c in df.columns if "totaleimpon" in c.lower())
        iva_col = next((c for c in df.columns if "totaleconiva" in c.replace(" ","").lower()), None)
        tot = next(c for c in df.columns if c.lower().strip()=="totale" or "totale" in c.lower())

        df[imp] = coerce_numeric(df[imp])
        if iva_col:
            df[iva_col] = coerce_numeric(df[iva_col])
        df[tot] = coerce_numeric(df[tot])

        base = prest

        # ========== AUTO APPRENDIMENTO PASS ==========
        learned = {norm(k) for k in (mem | new).keys()}
        df["_clean"] = df[base].astype(str).map(norm)
        candidates = sorted([t for t in df["_clean"].unique() if t not in learned])

        auto_added_now = []
        if model_B and vectorizer_B and candidates:
            X = vectorizer_B.transform(candidates)
            probs = model_B.predict_proba(X)
            preds = model_B.classes_[probs.argmax(axis=1)]
            confs = probs.max(axis=1)
            for t, p, c in zip(candidates, preds, confs):
                if float(c) >= auto_thresh:
                    new[t] = p
                    auto_added_now.append((t, p, float(c)))

        if auto_added_now:
            st.session_state.new = new
            st.session_state.auto_added.extend(auto_added_now)

        # Classify rows (using classify_B that includes memory & rules)
        df["CategoriaFinale"] = df[prest].apply(lambda x: classify_B(x, mem | new))

    # Remove Privato / Professionista
    df = df[~df["CategoriaFinale"].str.lower().isin(["privato","professionista"])]

    # ===== SHOW AUTO-LEARNED THIS RUN =====
    if st.session_state.auto_added:
        with st.expander(f"🤖 Suggerimenti AI da confermare ({len(st.session_state.auto_added)})"):
            auto_df = pd.DataFrame(st.session_state.auto_added, columns=["Termine", "Categoria", "Confidenza"])
            auto_df = auto_df.sort_values("Confidenza", ascending=False)
            st.dataframe(auto_df, use_container_width=True)

        if st.button("✅ Conferma questi suggerimenti"):
            for t, p, c in st.session_state.auto_added:
                new[t] = p
            st.session_state.mem.update(new)
            github_save_json(GITHUB_FILE_A if mode=="A" else GITHUB_FILE_B, st.session_state.mem)
            st.session_state.auto_added = []
            st.session_state.new = {}
            st.success("✅ Suggerimenti confermati e salvati.")

    # ===== LEARNING INTERFACE (manuale per ciò che resta) =====
    learned = {norm(k) for k in (mem | new).keys()}
    pending = [t for t in sorted(df["_clean"].unique()) if t not in learned]

    if pending:
        idx = st.session_state.idx
        if idx >= len(pending):
            idx = 0
            st.session_state.idx = 0
        term = pending[idx]
        opts = list(RULES_A.keys()) if mode=="A" else list(RULES_B.keys())
        last = st.session_state.get("last_cat", opts[0])
        default_index = opts.index(last) if last in opts else 0

        st.warning(f"🧠 Da classificare {idx+1}/{len(pending)} → “{term}”")
        cat_sel = st.selectbox("Categoria:", opts, index=default_index)

        if st.button("✅ Salva e prossimo"):
            new[norm(term)] = cat_sel
            st.session_state.new = new
            st.session_state.last_cat = cat_sel

            if idx + 1 >= len(pending):
                # Fine: salva su GitHub
                mem.update(new)
                github_save_json(GITHUB_FILE_A if mode=="A" else GITHUB_FILE_B, mem)
                st.session_state.mem = mem
                st.session_state.new = {}
                st.session_state.idx = 0
                st.success("🎉 Tutto classificato e salvato su GitHub!")
                st.rerun()

            st.session_state.idx = idx + 1
            st.rerun()

        st.stop()

    # ===== REPORT =====
    df = df.drop(columns=["_clean"], errors="ignore")

    if mode == "A":
        # DrVeto: Quantità(%) e Netto (dopo sconto)
        studio = df.groupby("CategoriaFinale").agg({qta:"sum", netto:"sum"}).reset_index()
        studio.columns = ["Categoria","Qtà","Netto"]
        studio["% Qtà"] = round_pct(studio["Qtà"])
        studio["% Netto"] = round_pct(studio["Netto"])
        studio.loc[len(studio)] = ["Totale", studio["Qtà"].sum(), studio["Netto"].sum(), 100, 100]
        ycol = "Netto"
        title = "Somma Netto per Categoria"

    else:
        # VetsGo: Imponibile + ConIVA, % su ConIVA
        studio = df.groupby("CategoriaFinale").agg({imp:"sum", iva_col:"sum"}).reset_index()
        studio.columns = ["Categoria","TotaleImponibile","TotaleConIVA"]
        studio["% Totale"] = round_pct(studio["TotaleConIVA"])
        studio.loc[len(studio)] = ["Totale", studio["TotaleImponibile"].sum(), studio["TotaleConIVA"].sum(), 100]
        studio["Categoria"] = pd.Categorical(studio["Categoria"], categories=ORDER_B, ordered=True)
        studio = studio.sort_values("Categoria")
        ycol = "TotaleConIVA"
        title = "Somma Totale con IVA per Categoria"

    st.dataframe(studio, use_container_width=True)

    fig, ax = plt.subplots(figsize=(8,5))
    ax.bar(studio["Categoria"], studio[ycol], color="steelblue")
    ax.set_title(title)
    plt.xticks(rotation=45, ha="right")
    buf = BytesIO(); plt.tight_layout(); plt.savefig(buf, format="png"); buf.seek(0)
    st.image(buf)

    wb = Workbook(); ws = wb.active; ws.title = "Report"
    for r in dataframe_to_rows(studio, index=False, header=True):
        ws.append(r)
    ws.add_image(XLImage(buf), f"A{len(studio)+3}")
    out = BytesIO(); wb.save(out)

    st.download_button("⬇️ Scarica Excel", data=out.getvalue(), file_name="StudioISA.xlsx")

    # ===== DASHBOARD =====
    if page == "Dashboard Annuale":
        st.header("📈 Dashboard Andamento Annuale")

        date_col = next(c for c in df.columns if "data" in c.replace(" ", "").lower())
        df[date_col] = (
            df[date_col].astype(str)
            .str.extract(r'(\d{1,4}[-/]\d{1,2}[-/]\d{2,4})')[0]
            .apply(lambda x: pd.to_datetime(x, dayfirst=True, errors="coerce"))
        )

        value_col = netto if mode=="A" else tot
        df["Anno"] = df[date_col].dt.year
        df["Mese"] = df[date_col].dt.to_period("M").astype(str)

        anni = sorted(df["Anno"].dropna().unique())
        if not anni:
            st.info("Nessuna data valida trovata per la dashboard.")
            st.stop()

        anno_sel = st.selectbox("Seleziona Anno:", anni, index=len(anni)-1)
        dfY = df[df["Anno"] == anno_sel]

        monthly = dfY.groupby("Mese")[value_col].sum().reset_index()
        all_months = pd.period_range(f"{anno_sel}-01", f"{anno_sel}-12", freq="M").astype(str)
        monthly = monthly.set_index("Mese").reindex(all_months, fill_value=0).reset_index().rename(columns={"index":"Mese"})

        st.subheader("Trend Fatturato Mensile")
        st.line_chart(monthly.set_index("Mese"))

        catshare = dfY.groupby("CategoriaFinale")[value_col].sum().reset_index()
        catshare["%"] = round_pct(catshare[value_col])
        st.subheader("Ripartizione per Categoria")
        st.bar_chart(catshare.set_index("CategoriaFinale")["%"])

        area = dfY.groupby(["Mese", "CategoriaFinale"])[value_col].sum().reset_index()
        area = area.pivot(index="Mese", columns="CategoriaFinale", values=value_col).fillna(0)
        st.subheader("Andamento Categorie nel Tempo")
        st.area_chart(area)

    # =============== REGISTRO IVA ===========
def add_simple_field(p, instr):
    r = p.add_run()
    for t, text in (('begin', ''), ('instrText', instr), ('separate', ''), ('end', '')):
        el = OxmlElement('w:fldChar') if t != 'instrText' else OxmlElement('w:instrText')
        if t != 'instrText':
            el.set(qn('w:fldCharType'), t)
        else:
            el.set(qn('xml:space'), 'preserve'); el.text = text
        r._r.append(el)

def add_lastpage_field(p, start_at):
    r = p.add_run()
    # { =
    el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), 'begin'); r._r.append(el)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = '= '; r._r.append(it)
    # { NUMPAGES }
    e1 = OxmlElement('w:fldChar'); e1.set(qn('w:fldCharType'), 'begin'); r._r.append(e1)
    it2 = OxmlElement('w:instrText'); it2.set(qn('xml:space'), 'preserve'); it2.text = 'NUMPAGES'; r._r.append(it2)
    e2 = OxmlElement('w:fldChar'); e2.set(qn('w:fldCharType'), 'end'); r._r.append(e2)
    # + offset -1
    it3 = OxmlElement('w:instrText'); it3.set(qn('xml:space'), 'preserve')
    it3.text = f' + {int(start_at)} - 1 '; r._r.append(it3)
    e3 = OxmlElement('w:fldChar'); e3.set(qn('w:fldCharType'), 'separate'); r._r.append(e3)
    p.add_run('0')
    e4 = OxmlElement('w:fldChar'); e4.set(qn('w:fldCharType'), 'end'); p._p.append(e4)

def add_field_run(paragraph, field):
    r = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    
def render_registro_iva():
    if st.sidebar.button("🔓 Logout"):
        st.session_state.pop("logged_user", None)
        st.rerun()

    user = load_users().get(logged_user, {})
    permissions = user.get("permissions", {})
    is_admin = user.get("role") == "admin"
    user_clinics = user.get("clinics", [])

    # ✅ Se non può usare il Registro IVA → STOP
    if not permissions.get("use_registro_iva", False):
        st.error("⛔ Non hai permesso di usare il Registro IVA.")
        st.stop()

    st.header("📄 Registro IVA")

    config_all = load_clinic_config()

    # ✅ Cliniche visibili
    if is_admin:
        cliniche_disponibili = list(config_all.keys())
    else:
        cliniche_disponibili = user_clinics

        if not cliniche_disponibili:
            st.warning("⚠️ Nessuna clinica assegnata. Contatta un amministratore.")
            st.stop()

    # ✅ Funzione che determina se l'utente può modificare quella clinica
    def can_edit_this_clinic(clinic_name):
        if is_admin:
            return True
        return permissions.get("manage_clinics", False) and clinic_name in user_clinics

    # ✅ "+ Nuova Clinica" mostrata solo se admin o può gestire cliniche
    can_add_clinic = is_admin or permissions.get("manage_clinics", False)
    choices = (["+ Nuova Clinica"] if can_add_clinic else []) + cliniche_disponibili

    clinica_scelta = st.selectbox("Seleziona Clinica", choices)


    # ➕ Aggiunta nuova clinica
        # ➕ Aggiunta nuova clinica
    if clinica_scelta == "+ Nuova Clinica":
        if not can_add_clinic:
            st.warning("🔒 Non hai permesso di aggiungere cliniche.")
            st.stop()

        struttura = st.text_input("Nome Struttura (Nuova)")
        via_ui = st.text_input("Via")
        cap_ui = st.text_input("CAP")
        citta_ui = st.text_input("Città")
        provincia_ui = st.text_input("Provincia (sigla)", max_chars=2)
        piva = st.text_input("Partita IVA")
        pagina_iniziale = st.number_input("Numero pagina iniziale", min_value=1, max_value=999, value=1)

        if st.button("➕ Aggiungi Clinica"):
            config_all[struttura] = {
                "struttura": struttura,
                "via": via_ui,
                "cap": cap_ui,
                "citta": citta_ui,
                "provincia": provincia_ui.upper(),
                "piva": piva,
                "pagina_iniziale_default": int(pagina_iniziale)
            }
            save_clinic_config(config_all)
            st.success("✅ Nuova clinica aggiunta!")
            st.rerun()

    # ✏️ Modifica clinica esistente
    else:
        cfg = config_all[clinica_scelta]
        readonly = not can_edit_this_clinic(clinica_scelta)

        struttura = st.text_input("Nome Struttura", cfg.get("struttura",""), disabled=readonly)
        via_ui = st.text_input("Via", cfg.get("via",""), disabled=readonly)
        cap_ui = st.text_input("CAP", cfg.get("cap",""), disabled=readonly)
        citta_ui = st.text_input("Città", cfg.get("citta",""), disabled=readonly)
        provincia_ui = st.text_input("Provincia (sigla)", max_chars=2, value=cfg.get("provincia",""), disabled=readonly)
        piva = st.text_input("Partita IVA", cfg.get("piva",""), disabled=readonly)
        pagina_iniziale = st.number_input("Numero pagina iniziale", min_value=1, max_value=999,
                                          value=cfg.get("pagina_iniziale_default",1))

        if not readonly and st.button("💾 Salva modifiche"):
            config_all[clinica_scelta] = {
                "struttura": struttura,
                "via": via_ui,
                "cap": cap_ui,
                "citta": citta_ui,
                "provincia": provincia_ui.upper(),
                "piva": piva,
                "pagina_iniziale_default": int(pagina_iniziale)
            }
            save_clinic_config(config_all)
            st.success("✅ Modifiche salvate!")
            st.rerun()

    # 📂 Upload file
    file = st.file_uploader("Carica il file Registro IVA (Excel)", type=["xlsx", "xls"])
    if not file or not struttura:
        return

    # --- Leggi Excel ---
    df_raw = pd.read_excel(file)

    # --- Normalizza i nomi colonna (spazi multipli, NBSP, trim) ---
    def norm_col(c: str) -> str:
        c = str(c).replace("\u00A0", " ")
        c = re.sub(r"\s+", " ", c).strip()
        return c

    df_raw.columns = [norm_col(c) for c in df_raw.columns]

    # --- Helper: trova colonna per nome "normalizzato" (case-insensitive) ---
    def find_col_by_norm(df, target: str):
        target_n = norm_col(target).lower()
        # 1) match esatto case-insensitive
        for c in df.columns:
            if norm_col(c).lower() == target_n:
                return c
        # 2) match che inizia per (es. "Totale imponibile ..." vs "Totale imponibile")
        for c in df.columns:
            if norm_col(c).lower().startswith(target_n):
                return c
        # 3) match contenuto (ultima spiaggia)
        for c in df.columns:
            if target_n in norm_col(c).lower():
                return c
        return None

    # Mappa numerica interna (alias)   <nome logico cercato -> alias interno>
    wanted_num = {
        "Totale Netto": "tot_netto",
        "Totale ENPAV": "tot_enpav",
        "Totale imponibile": "tot_imponibile",   # <- robusto: minuscole/maiuscole non contano
        "Totale IVA": "tot_iva",
        "Totale Sconto": "tot_sconto",
        "Rit. d'acconto": "tot_rit",
        "Totale": "totale",
    }

    # Trova i nomi REALI presenti nel file per ciascun “wanted”
    real_num_cols = {}
    for wanted_name, alias in wanted_num.items():
        col_found = find_col_by_norm(df_raw, wanted_name)
        if col_found is not None:
            real_num_cols[alias] = col_found

    # Prepara serie numeriche per i totali (senza duplicare le colonne stampate)
    df_num = pd.DataFrame(index=df_raw.index)
    for alias, real_col in real_num_cols.items():
        s = df_raw[real_col]
        # conversione robusta a numerico
        if pd.api.types.is_numeric_dtype(s):
            df_num[alias] = pd.to_numeric(s, errors="coerce").fillna(0)
        else:
            df_num[alias] = (
                s.astype(str)
                 .str.replace(r"\s", "", regex=True)
                 .str.replace("€", "", regex=False)
                 .str.replace(".", "", regex=False)
                 .str.replace(",", ".", regex=False)
                 .pipe(pd.to_numeric, errors="coerce")
                 .fillna(0)
            )

    # --- Colonne da mostrare a video / esportare (usiamo i nomi REALI del file) ---
    # Mantieni ordine e verifica presenza
    preferred_display = [
        "Data", "Numero", "Cliente", "P. IVA", "Codice Fiscale",
        "Indirizzo", "CAP", "Città",
        "Totale Netto", "Totale ENPAV", "Totale imponibile",  # <- qui il nome richiesto dal file
        "Totale IVA", "Totale Sconto", "Rit. d'acconto", "Totale",
    ]
    # Per "Totale imponibile" usa la versione reale trovata (potrebbe avere variante)
    real_tot_imp = find_col_by_norm(df_raw, "Totale imponibile")
    if real_tot_imp and real_tot_imp not in preferred_display:
        # sostituisci l’etichetta di comodo con quella reale
        preferred_display = [real_tot_imp if x.lower() == "totale imponibile" else x for x in preferred_display]

    cols_presenti = [c for c in preferred_display if c in df_raw.columns]
    if not cols_presenti:
        st.error("❌ Il file non contiene le colonne richieste per il Registro IVA.")
        return

    df_display = df_raw.loc[:, cols_presenti].copy()
    df_display = df_display.iloc[::-1].reset_index(drop=True)

    # --- Aggiungi provincia nella colonna "Città" ---
    has_prov_in_file = find_col_by_norm(df_raw, "Provincia")

    if "Città" in df_display.columns:
        citta_clean = df_display["Città"].astype(str).str.strip()

        if has_prov_in_file:
            prov_col = find_col_by_norm(df_raw, "Provincia")
            prov = df_raw[prov_col].astype(str).str.strip().str.upper()
            df_display["Città"] = citta_clean + " (" + prov + ")"

        else:
            # Provincia NON nel file → usa quella inserita dall’utente, se presente
            if provincia_ui:
                df_display["Città"] = citta_clean + " (" + provincia_ui.upper().strip() + ")"
            else:
                df_display["Città"] = citta_clean

    # CAP pulito (evita “40.033,00” o simili)
    if "CAP" in df_display.columns:
        df_display["CAP"] = (
            df_display["CAP"].astype(str)
            .str.replace(r"[^\dA-Za-z]", "", regex=True)
        )

    # Intervallo date + anno
    if "Data" in df_display.columns:
        ds = pd.to_datetime(df_display["Data"], dayfirst=True, errors="coerce")
        data_min = ds.min()
        data_max = ds.max()
        anno = int(ds.dt.year.dropna().mode()[0]) if ds.notna().any() else datetime.now().year
        data_min_str = data_min.strftime("%d/%m/%Y") if pd.notna(data_min) else "-"
        data_max_str = data_max.strftime("%d/%m/%Y") if pd.notna(data_max) else "-"
    else:
        anno = datetime.now().year
        data_min_str = "-"
        data_max_str = "-"

    # Dati indirizzo (UI ha priorità; se vuoti, prova a leggere la prima riga del file)
    via_file = str(df_display["Indirizzo"].iloc[0]) if "Indirizzo" in df_display.columns and not df_display.empty else ""
    cap_file = str(df_display["CAP"].iloc[0]) if "CAP" in df_display.columns and not df_display.empty else ""
    citta_file = str(df_display["Città"].iloc[0]) if "Città" in df_display.columns and not df_display.empty else ""

    via = via_ui or via_file or ""
    cap_print = cap_ui or cap_file or ""
    citta_print = citta_ui or citta_file or ""
    if provincia_ui:
        citta_print = f"{citta_print} ({provincia_ui.upper()})".strip()

    # Mostra anteprima tabella
    st.dataframe(df_display, use_container_width=True)
        # --- CONTROLLO ANOMALIE ----------------------------------------------------
    warnings = []

    # 1) IVA zero ma imponibile > 0
    if "Totale imponibile" in df_display.columns and "Totale IVA" in df_display.columns:
        imp = pd.to_numeric(df_display["Totale imponibile"], errors="coerce").fillna(0)
        iva = pd.to_numeric(df_display["Totale IVA"], errors="coerce").fillna(0)
        mask = (imp > 0) & (iva == 0)
        if mask.any():
            warnings.append(f"➡️ {mask.sum()} righe hanno imponibile > 0 ma IVA = 0 (possibile aliquota mancante).")

    # 2) Totale minore di imponibile + iva
    if "Totale" in df_display.columns:
        tot = pd.to_numeric(df_display["Totale"], errors="coerce").fillna(0)
        calc = imp + iva
        mask2 = tot < (calc - 0.01)  # tolleranza 1 cent
        if mask2.any():
            warnings.append(f"➡️ {mask2.sum()} righe hanno Totale più basso della somma Imponibile + IVA.")

    # 3) PIVA vuota → possibile privato (solo informativo)
    if "P. IVA" in df_display.columns:
        mask3 = df_display["P. IVA"].astype(str).str.strip().eq("")
        if mask3.any():
            warnings.append(f"ℹ️ {mask3.sum()} righe senza Partita IVA (probabili privati).")

    # 4) CAP non numerico
    if "CAP" in df_display.columns:
        cap_wrong = ~df_display["CAP"].astype(str).str.match(r"^\d{4,5}$")
        if cap_wrong.any():
            warnings.append(f"⚠️ {cap_wrong.sum()} CAP non hanno un formato valido (es: 40121).")

    # 5) Date mancanti o non valide
    if "Data" in df_display.columns:
        dt = pd.to_datetime(df_display["Data"], dayfirst=True, errors="coerce")
        if dt.isna().any():
            warnings.append(f"⚠️ {dt.isna().sum()} righe hanno una data non valida.")

    # ---- Mostra riepilogo ------------------------------------------------------
    if warnings:
        st.warning(
            "⚠️ **Controllo Anomalie nel Registro IVA**:\n\n" +
            "\n".join(warnings) +
            "\n\nPuoi comunque procedere con la generazione del file."
        )
    else:
        st.success("✅ Nessuna anomalia rilevata.")

    if not st.button("🧾 Genera Registro IVA (DOCX)"):
        return

    with st.spinner("Generazione del Registro IVA in corso..."):
        # Versione stringa per scrittura veloce in Word
        df_display_str = df_display.fillna("").astype(str)

        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn

        doc = Document()

        # Layout orizzontale + margini
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_number_start = pagina_iniziale
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)

        pgNumType.set(qn('w:start'), str(int(pagina_iniziale)))

        settings = doc.settings._element
        upd = OxmlElement('w:updateFields'); upd.set(qn('w:val'), 'true')
        settings.append(upd)

        section.header.is_linked_to_previous = False

        # Stile base
        style = doc.styles["Normal"]
        style.font.name = "Aptos Narrow"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Narrow")
        style.font.size = Pt(8)

        # Header
        header = section.header
        hdr_table = header.add_table(rows=1, cols=2, width=Inches(11.0))
        hdr_left, hdr_right = hdr_table.rows[0].cells

        # Sinistra
        pL = hdr_left.paragraphs[0]
        pL.alignment = WD_ALIGN_PARAGRAPH.LEFT

        r1 = pL.add_run(struttura + "\n")
        r1.font.name = "Segoe UI"; r1.font.size = Pt(14)

        r2 = pL.add_run(" ".join(x for x in [via, cap_print, citta_print] if x) + "\n")
        r2.font.name = "Segoe UI"; r2.font.size = Pt(12)

        r3 = pL.add_run(f"P.IVA {piva}")
        r3.font.name = "Aptos Narrow"; r3._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Narrow")
        r3.font.size = Pt(10); r3.bold = True

        # Destra
        pR = hdr_right.paragraphs[0]
        pR.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        r4 = pR.add_run(f"ANNO {anno}\n")
        r4.font.name = "Calibri"; r4.font.size = Pt(10)

        r5 = pR.add_run(f"Entrate dal {data_min_str} al {data_max_str}\n")
        r5.font.name = "Aptos Narrow"; r5._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Narrow")
        r5.font.size = Pt(10); r5.bold = True

        p_page = hdr_right.add_paragraph()
        p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        r = p_page.add_run("Pag. "); r.font.name="Aptos Narrow"; r.font.size=Pt(10)
        add_simple_field(p_page, "PAGE")
        r2 = p_page.add_run(" di "); r2.font.name="Aptos Narrow"; r2.font.size=Pt(10)
        add_lastpage_field(p_page, pagina_iniziale)

        doc.add_paragraph()

                # Tabella
        rows, cols = df_display_str.shape
        table = doc.add_table(rows=rows + 1, cols=cols)
        table.style = "Table Grid"

        # Intestazioni
        for j, col_name in enumerate(df_display_str.columns):
            cell = table.cell(0, j)
            p = cell.paragraphs[0]
            run = p.add_run(col_name)
            run.bold = True

        # Dati
        for i in range(rows):
            row_vals = df_display_str.iloc[i]
            row_cells = table.rows[i + 1].cells
            for j in range(cols):
                row_cells[j].text = row_vals.iloc[j]

        doc.add_paragraph()

        # Totali finali (dalla vista numerica)
        def euro_it(v: float) -> str:
            s = f"{v:,.2f}"
            return s.replace(",", "X").replace(".", ",").replace("X", ".")

        tot_netto = df_num.get("tot_netto", pd.Series([], dtype=float)).sum()
        tot_enpav = df_num.get("tot_enpav", pd.Series([], dtype=float)).sum()
        tot_imp   = df_num.get("tot_imponibile", pd.Series([], dtype=float)).sum()
        tot_iva   = df_num.get("tot_iva", pd.Series([], dtype=float)).sum()
        tot_sco   = df_num.get("tot_sconto", pd.Series([], dtype=float)).sum()
        tot_rit   = df_num.get("tot_rit", pd.Series([], dtype=float)).sum()
        tot_tot   = df_num.get("totale", pd.Series([], dtype=float)).sum()

        # Scrittura totali in grassetto
        def add_bold_total(label, value):
            p = doc.add_paragraph()
            r1 = p.add_run(f"{label}: ")
            r1.bold = True
            r2 = p.add_run(f"{euro_it(value)} €")
            r2.bold = True

        add_bold_total("Totale Netto", tot_netto)
        add_bold_total("Totale ENPAV", tot_enpav)
        add_bold_total("Totale Imponibile", tot_imp)
        add_bold_total("Totale IVA", tot_iva)
        add_bold_total("Totale Sconto", tot_sco)
        add_bold_total("Ritenuta d'acconto", tot_rit)
        add_bold_total("Totale complessivo", tot_tot)

        # Esporta DOCX
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc_bytes = buf.getvalue()

    # ---- FUORI DALLO SPINNER ----
    st.success("✅ Registro IVA generato.")
    st.download_button(
        "⬇️ Scarica Registro IVA (Word)",
        data=doc_bytes,
        file_name=f"Registro_IVA_{anno}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )



if __name__ == "__main__":
    if page == "📄 Registro IVA":
        render_registro_iva()
    elif page == "👤 Gestione Utenti":
        render_user_management()
    else:
        main()


































































